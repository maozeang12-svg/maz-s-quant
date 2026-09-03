from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

output_path = r"C:\Users\11\Documents\Kimi\Workspaces\个人提升计划\毛泽昂-简历-量化研究员.docx"

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

def set_run_font(run, name_en="Calibri", name_cn="微软雅黑", size=10.5, bold=False, color=None):
    run.font.name = name_en
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), name_cn)

def add_heading(doc, text, size=14, bold=True, color=RGBColor(0x2C, 0x3E, 0x50), space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p

def add_bullet(doc, text, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
    run = p.add_run(text)
    set_run_font(run, size=10)
    return p

def add_normal_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        set_run_font(run_bold, size=10, bold=True)
        run_text = p.add_run(text)
        set_run_font(run_text, size=10)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10)
    return p

# ========== 标题区域 ==========
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(6)
title_run = title_p.add_run("毛泽昂")
set_run_font(title_run, size=22, bold=True, color=RGBColor(0x2C, 0x3E, 0x50))

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_p.paragraph_format.space_after = Pt(12)
info_run = info_p.add_run("男 | 27岁 | 杭州 | 199****6131 | 2660134803@qq.com")
set_run_font(info_run, size=10, color=RGBColor(0x5D, 0x6D, 0x7E))

intent_p = doc.add_paragraph()
intent_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
intent_p.paragraph_format.space_after = Pt(12)
intent_run = intent_p.add_run("求职意向：量化研究员 / 量化策略分析师  |  期望薪资：14-20K  |  期望城市：杭州")
set_run_font(intent_run, size=10, color=RGBColor(0x5D, 0x6D, 0x7E))

# ========== 分隔线 ==========
sep = doc.add_paragraph()
sep.paragraph_format.space_before = Pt(0)
sep.paragraph_format.space_after = Pt(6)
sep_run = sep.add_run("─" * 60)
set_run_font(sep_run, size=8, color=RGBColor(0xD5, 0xD8, 0xDC))

# ========== 个人优势 ==========
add_heading(doc, "个人优势")

advantages = [
    "4年实盘交易经验，熟悉A股/期货/美股跨市场交易，具备真实风控体感与市场微观结构认知",
    "现任量化研究员，具备「研究假设→数据分析→回测验证→实盘归因」的完整策略研发闭环能力",
    "深耕A股短线策略：集合竞价选股、连板放量突破、MACD水上金叉监控，具备可验证的回测样本与归因报告",
    "独立搭建量化基础设施：行情数据库（亿级数据迁移）、Python回测框架、盘中实时监控平台",
    "扎实的Python数据处理能力（pandas/numpy/matplotlib/akshare），熟悉MySQL/SQLite数据管理与SQL查询",
]
for adv in advantages:
    add_bullet(doc, adv)

# ========== 技术栈 ==========
add_heading(doc, "技术栈")

tech_items = [
    ("编程与数据分析：", "Python（pandas / numpy / matplotlib / akshare）| Pine Script（TradingView）"),
    ("数据库：", "MySQL | SQLite | PostgreSQL | SQL（熟练查询与优化）"),
    ("数据工具：", "Tushare | AkShare | 米筐量化平台 | 通达信公式语言 | TQ接口"),
    ("工程化：", "Docker | Flask | APScheduler | SQLAlchemy"),
    ("策略方法：", "多因子选股 | 量价分析 | 集合竞价 | 趋势跟踪 | 风险归因 | 样本内外验证"),
]
for prefix, content in tech_items:
    add_normal_para(doc, content, bold_prefix=prefix)

# ========== 工作经历 ==========
add_heading(doc, "工作经历")

# 沐龙实业
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(2)
run1 = p.add_run("浙江沐龙实业集团有限责任公司")
set_run_font(run1, size=11, bold=True)
run2 = p.add_run("  |  量化研究员  |  2025.05 - 至今")
set_run_font(run2, size=10)

ml_work = [
    ("A股短线策略研发与回测：", "集合竞价策略（米筐平台）：连板过滤 + 集合竞价量异常（倍量≥2倍/缩量≤0.5倍且高开>5%），构建可复现的短线选股信号"),
    ("", "放量突破策略：ADX>20 + 成交量>1.2倍20日均量 + 站上20日均线 + 突破60日新高的四条件复合选股，全池未触发时将「空仓等待」作为有效结论写入交易计划"),
    ("", "季节性回测：按月（7/8/9月）回测大盘及板块指数胜率，验证日线重算与月线法一致性（误差<0.01%），形成不同市场环境下的择时决策依据"),
    ("", "T+0 ETF波段策略：构建ETF交易池（港股ETF/香港证券/恒生科技），设定同赛道合并仓位上限30%的集中度约束；三年日线代理回测筛选最优标的"),
    ("Python回测框架：", "基于akshare+pandas+matplotlib搭建本地回测框架，完成策略从Pine Script到Python的迁移验证"),
    ("", "跨品种验证：铜期货日线vs60分钟级别测试，识别策略在周期性商品上失效（夏普约-0.20、最大回撤约-80%），判定60分钟级别短期盈利为样本量不足的统计假象，形成清晰的策略适用边界认知"),
    ("", "回测计入万3.5手续费+2 tick滑点，输出可落地的实盘策略参数方案"),
    ("数据处理与因子构建：", "搭建A股量化行情数据库：设计14张表覆盖日线/分钟/Tick/财务三表/复权因子/股票基础信息/交易日历，主导存储选型演进路线"),
    ("", "完成1.97亿行数据迁移（195分钟），发现3~7月覆盖仅58%的漏股问题并定位根因，制定数据口径规则（统一代码格式、复权分离、Tick主键防丢单）"),
    ("", "自编集合竞价量能指标：将同花顺公式改写为通达信指标，绘制今昨竞价量对比柱，解决历史竞价数据缺失问题"),
    ("风控与归因体系：", "自研趋势跟踪策略（Pine Script）：RSI双阈值（30/50）入场 + WMA动态止损 + 硬止损（买入价×95%），基于520笔回测样本数据化推导风控规则"),
    ("", "60分钟时间止损（≤60分钟胜率82% vs 超时31%）、14:55强制清仓禁隔夜（日内胜率68% vs 隔夜50%）、开盘14分钟噪声段过滤、单日熔断"),
    ("", "制度化复盘：撰写周度投资周报与月度投资总结，8月复盘后新增三条硬约束——浮亏不加仓、卖出当日不追高回补、月末利润保护降仓"),
    ("", "诚实记录策略局限：未扣手续费/滑点、一字板开板坑、8月数据断层、样本量偏小，体现对过拟合与数据偏差的认知"),
    ("跨市场交易：", "美股：以SOXL（三倍做多半导体ETF）为主要标的，执行「恐慌性下跌抄底+反弹兑现」短线策略；跟踪SMCI等个股财报事件"),
    ("", "加密资产：以ETH为主要标的，回调分批做多，不碰山寨币"),
    ("", "宏观研究：跟踪美债收益率曲线倒挂、降息预期、原油价格与对冲基金仓位对风险资产的传导路径"),
    ("监控工具：", "通达信选股公式体系：创业板强势股筛选、MACD水上复合筛选、水上死叉选股，多条件逐步迭代"),
    ("", "TQ Python盘中实时监控：对接通达信TQ接口，监控约100只自选股30/60分钟MACD金叉信号，触发记录本地存档（标的、价格、时间、触发级别），实现盘中自动预警"),
]
for prefix, content in ml_work:
    add_normal_para(doc, content, bold_prefix=prefix)

# 海亮
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(2)
run1 = p.add_run("浙江海亮股份有限公司")
set_run_font(run1, size=11, bold=True)
run2 = p.add_run("  |  证券交易员  |  2024.04 - 2025.01")
set_run_font(run2, size=10)

hl_work = [
    "负责SHFE/LME/COMEX等市场的套保套利交易执行，策略包括跨月套利、跨市套利、期现套利",
    "每日晨评撰写，对行情趋势及风险敞口进行预判；管理期货及现货头寸，确保头寸与公司风险管理政策相符",
    "及时响应突发情况，将敞口控制在可控范围内；优化成交点位，熟练掌握国内外各类交易软件",
]
for item in hl_work:
    add_bullet(doc, item)

# 秉越
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(2)
run1 = p.add_run("浙江秉越资产管理有限公司")
set_run_font(run1, size=11, bold=True)
run2 = p.add_run("  |  证券交易员  |  2022.11 - 2024.01")
set_run_font(run2, size=10)

by_work = [
    "执行期货/股票交易指令，13个月零错误；通过量化回测改进交易模型，完善个人期货交易体系",
    "提示大盘风险，帮助大账户减少300万以上亏损；9个月测试周期：最大回撤9%，收益率100%",
    "有较好的股票日内交易能力，帮助公司部分股票持仓成本降低5%",
]
for item in by_work:
    add_bullet(doc, item)

# 亿城
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(2)
run1 = p.add_run("杭州亿城实业有限公司")
set_run_font(run1, size=11, bold=True)
run2 = p.add_run("  |  证券交易员  |  2022.07 - 2022.12")
set_run_font(run2, size=10)
add_bullet(doc, "为期1月模拟盘业绩50%收益率，4个月实盘交易20%收益率")

# 弈宸
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(2)
run1 = p.add_run("杭州弈宸私募基金管理有限公司")
set_run_font(run1, size=11, bold=True)
run2 = p.add_run("  |  股票交易员  |  2022.02 - 2022.07")
set_run_font(run2, size=10)
add_bullet(doc, "每日寻找T0交易机会，训练交易反应能力与止损执行力")

# ========== 个人项目 ==========
add_heading(doc, "个人项目（业余时间）")

projects = [
    ("A股量化行情数据库", "设计14张表覆盖日线/分钟/Tick/财务三表/复权因子/股票基础信息/交易日历；主导存储选型演进：SQLite→PostgreSQL+TimescaleDB→时序库；完成1.97亿行数据迁移（195分钟），发现3~7月覆盖仅58%的漏股问题并定位根因；工具：MySQL | Python | pandas"),
    ("连板竞价策略回测", "定义策略假设：连续涨停+竞价放量→次日超额收益；主导加入「情绪闸门(ADJS≥0)」过滤条件，组合胜率从全样本50.7%提升至66.7%（33笔）；完成分月/分开盘涨幅/分情绪的归因分析；诚实记录局限：未扣手续费/滑点、8月数据断层、样本量偏小；工具：Python | pandas | MySQL"),
    ("股票监控报警平台", "搭建6层架构监控平台：DataFeed→策略→Engine→Notifier→Storage→Web；行情源双源兜底（Tushare主+AkShare备），告警走钉钉+邮件；策略热加载机制，Docker容器化部署；工具：Python | Flask | SQLAlchemy | Docker | APScheduler"),
]
for title, desc in projects:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run_title = p.add_run(title)
    set_run_font(run_title, size=10, bold=True)
    run_desc = p.add_run("  |  " + desc)
    set_run_font(run_desc, size=10)

# ========== 教育经历 ==========
add_heading(doc, "教育经历")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run1 = p.add_run("杭州师范大学钱江学院")
set_run_font(run1, size=10.5, bold=True)
run2 = p.add_run("  |  经济学  |  本科  |  2020 - 2022")
set_run_font(run2, size=10)

edu = [
    "主修：宏观经济学、微观经济学、国际经济学、证券投资学、信托、经济法",
    "荣誉：财信期货高校杯二等奖、浙江省大学生投资理财大赛复赛、全国黄炎培杯理财大赛参与者",
]
for item in edu:
    add_bullet(doc, item)

# ========== 资格证书 ==========
add_heading(doc, "资格证书")

certs = doc.add_paragraph()
certs.paragraph_format.space_after = Pt(2)
run = certs.add_run("基金从业资格证  |  期货从业资格证")
set_run_font(run, size=10)

# 保存
doc.save(output_path)
print(f"简历已保存: {output_path}")

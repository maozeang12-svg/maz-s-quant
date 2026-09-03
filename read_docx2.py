from docx import Document
import os

path = r"C:\Users\11\Downloads\项目经历整理-金融分析与量化开发.docx"

doc = Document(path)

print("=== 文档内容 ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"{para.text}")

print("\n=== 表格内容 ===")
for t, table in enumerate(doc.tables):
    print(f"\n--- 表格 {t+1} ---")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            print(" | ".join(cells))
